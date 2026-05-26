"""
Document Service - Xử lý upload, trích xuất văn bản, phân mảnh (chunking) và tạo embedding.

Quy chuẩn chunking theo chunkdocs.md:
- chunk_size = 512 tokens (~1500 ký tự)
- chunk_overlap = 64 tokens (~150-200 ký tự)
- Mỗi định dạng file có chiến lược chunking riêng biệt.
"""
import hashlib
import logging
import os
import re
import uuid
from datetime import datetime, timezone
from io import BytesIO
from typing import Optional

import pandas as pd
import tiktoken
from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.document import Document, DocumentChunk, DocumentStatus
from app.services import embedding_service

logger = logging.getLogger(__name__)

# --- Cấu hình Chunking ---
PARENT_CHUNK_SIZE = 1024  # tokens
PARENT_CHUNK_OVERLAP = 128  # tokens
CHILD_CHUNK_SIZE = 256  # tokens
CHILD_CHUNK_OVERLAP = 32  # tokens
ENCODING = tiktoken.get_encoding("cl100k_base")

# Thư mục lưu file upload
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


# =============================================================================
# Utility Functions
# =============================================================================

def count_tokens(text: str) -> int:
    """Đếm số tokens của chuỗi văn bản bằng tiktoken."""
    return len(ENCODING.encode(text))


def compute_checksum(file_bytes: bytes) -> str:
    """Tính SHA-256 checksum cho file bytes."""
    return hashlib.sha256(file_bytes).hexdigest()


# =============================================================================
# Chiến lược Chunking theo định dạng
# =============================================================================

def _recursive_split(text: str, separators: list[str], max_tokens: int = PARENT_CHUNK_SIZE) -> list[str]:
    """
    Phân mảnh đệ quy (Recursive Character Splitting).
    Thử tách theo separator đầu tiên, nếu chunk vẫn > max_tokens thì thử separator tiếp theo.
    Fallback cuối cùng: cắt cứng tại khoảng trắng gần nhất.
    """
    if not text or count_tokens(text) <= max_tokens:
        return [text] if text.strip() else []

    sep = separators[0] if separators else ""
    remaining_seps = separators[1:] if len(separators) > 1 else []

    if sep and sep in text:
        parts = text.split(sep)
    else:
        # Fallback: cắt cứng tại khoảng trắng
        words = text.split(" ")
        chunks = []
        current = ""
        for word in words:
            test = f"{current} {word}".strip()
            if count_tokens(test) > max_tokens and current:
                chunks.append(current)
                current = word
            else:
                current = test
        if current:
            chunks.append(current)
        return chunks

    # Ghép các parts lại thành chunks <= max_tokens
    chunks = []
    current_chunk = ""
    for part in parts:
        candidate = f"{current_chunk}{sep}{part}" if current_chunk else part
        if count_tokens(candidate) <= max_tokens:
            current_chunk = candidate
        else:
            if current_chunk:
                chunks.append(current_chunk)
            # Nếu part đơn lẻ vẫn > max_tokens, đệ quy chia nhỏ
            if count_tokens(part) > max_tokens and remaining_seps:
                chunks.extend(_recursive_split(part, remaining_seps, max_tokens))
            elif count_tokens(part) > max_tokens:
                # Cắt cứng
                chunks.extend(_recursive_split(part, [], max_tokens))
            else:
                current_chunk = part
    if current_chunk:
        chunks.append(current_chunk)

    return [c for c in chunks if c.strip()]


def _add_overlap(chunks: list[str], overlap_tokens: int = CHILD_CHUNK_OVERLAP) -> list[str]:
    """Thêm overlap giữa các chunks liên tiếp."""
    if len(chunks) <= 1:
        return chunks

    result = [chunks[0]]
    for i in range(1, len(chunks)):
        prev_tokens = ENCODING.encode(chunks[i - 1])
        overlap_text = ENCODING.decode(prev_tokens[-overlap_tokens:]) if len(prev_tokens) > overlap_tokens else chunks[i - 1]
        result.append(f"{overlap_text} {chunks[i]}")
    return result


def split_parent_child(
    text: str,
    separators: list[str] = ["\n\n", "\n", ". ", "? ", "! ", " ", ""],
    parent_size: int = PARENT_CHUNK_SIZE,
    parent_overlap: int = PARENT_CHUNK_OVERLAP,
    child_size: int = CHILD_CHUNK_SIZE,
    child_overlap: int = CHILD_CHUNK_OVERLAP,
) -> list[dict]:
    """
    Phân tách văn bản thành cấu trúc Parent-Child:
    Trả về danh sách các dict, mỗi dict chứa parent_content, metadata và danh sách các children.
    """
    # 1. Tách văn bản thành các Parent chunks
    raw_parents = _recursive_split(text, separators, max_tokens=parent_size)
    parent_chunks = _add_overlap(raw_parents, overlap_tokens=parent_overlap)
    
    result = []
    for parent_text in parent_chunks:
        # 2. Tách mỗi Parent chunk thành các Child chunks
        raw_children = _recursive_split(parent_text, separators, max_tokens=child_size)
        child_chunks = _add_overlap(raw_children, overlap_tokens=child_overlap)
        
        result.append({
            "parent_content": parent_text,
            "children": child_chunks
        })
    return result


def chunk_plain_text_parent_child(text: str) -> list[dict]:
    """Phân mảnh Plain Text (.txt) theo cấu trúc Parent-Child."""
    separators = ["\n\n", "\n", ". ", "? ", "! ", " ", ""]
    parent_child_data = split_parent_child(text, separators)
    for item in parent_child_data:
        item["metadata"] = {"source_type": "txt"}
    return parent_child_data


def chunk_markdown_parent_child(text: str) -> list[dict]:
    """
    Phân mảnh Markdown (.md) theo cấu trúc Parent-Child:
    - Cấp 1 (Parent): Tách theo Heading. Nếu heading section > PARENT_CHUNK_SIZE -> Chia nhỏ đệ quy.
    - Cấp 2 (Child): Mỗi Parent sẽ được chia nhỏ thành các Child chunks (CHILD_CHUNK_SIZE).
    """
    # Lớp 1: Tách theo headings
    heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
    sections = []
    last_end = 0
    current_heading = ""

    for match in heading_pattern.finditer(text):
        if last_end < match.start():
            content = text[last_end:match.start()].strip()
            if content:
                sections.append({"heading": current_heading, "content": content})
        current_heading = match.group(2).strip()
        last_end = match.end()

    # Phần còn lại sau heading cuối
    remaining = text[last_end:].strip()
    if remaining:
        sections.append({"heading": current_heading, "content": remaining})

    parent_child_data = []
    for section in sections:
        content = section["content"]
        heading = section["heading"]

        # Chia thành parent chunks
        if count_tokens(content) <= PARENT_CHUNK_SIZE:
            parent_chunks = [content]
        else:
            raw_parents = _recursive_split(content, ["\n\n", "\n", ". ", " "], max_tokens=PARENT_CHUNK_SIZE)
            parent_chunks = _add_overlap(raw_parents, overlap_tokens=PARENT_CHUNK_OVERLAP)

        for parent_text in parent_chunks:
            # Chia từng parent thành child chunks
            if count_tokens(parent_text) <= CHILD_CHUNK_SIZE:
                child_chunks = [parent_text]
            else:
                raw_children = _recursive_split(parent_text, ["\n\n", "\n", ". ", " "], max_tokens=CHILD_CHUNK_SIZE)
                child_chunks = _add_overlap(raw_children, overlap_tokens=CHILD_CHUNK_OVERLAP)

            parent_child_data.append({
                "parent_content": parent_text,
                "metadata": {"source_type": "markdown", "heading": heading},
                "children": child_chunks
            })

    return parent_child_data


def chunk_docx_parent_child(file_bytes: bytes) -> list[dict]:
    """
    Chunk Word (.docx): Chuyển đổi sang Markdown trung gian, rồi áp dụng chunk_markdown_parent_child.
    """
    doc = DocxDocument(BytesIO(file_bytes))
    md_lines = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue

        if "Heading 1" in style_name:
            md_lines.append(f"# {text}")
        elif "Heading 2" in style_name:
            md_lines.append(f"## {text}")
        elif "Heading 3" in style_name:
            md_lines.append(f"### {text}")
        elif "List" in style_name:
            md_lines.append(f"- {text}")
        else:
            md_lines.append(text)

    # Xử lý bảng trong Word
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        md_lines.append("| " + " | ".join(headers) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
        md_lines.append("")

    markdown_text = "\n".join(md_lines)
    return chunk_markdown_parent_child(markdown_text)


def chunk_pdf_parent_child(file_bytes: bytes) -> tuple[list[dict], int]:
    """
    Chunk PDF (.pdf): Trích xuất text bằng pypdf, áp dụng split_parent_child.
    Trả về (parent_child_data, page_count).
    """
    reader = PdfReader(BytesIO(file_bytes))
    page_count = len(reader.pages)
    full_text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            full_text += page_text + "\n\n"

    separators = ["\n\n", "\n", ". ", "? ", "! ", " ", ""]
    parent_child_data = split_parent_child(full_text, separators)
    for item in parent_child_data:
        item["metadata"] = {"source_type": "pdf"}
    return parent_child_data, page_count


def chunk_excel_csv_parent_child(file_bytes: bytes, file_name: str) -> list[dict]:
    """
    Chunk Excel/CSV: Mỗi dòng = 1 Parent chunk.
    Do dòng ngắn nên Child chunk trùng luôn với Parent chunk.
    """
    if file_name.endswith(".csv"):
        df = pd.read_csv(BytesIO(file_bytes))
    else:
        df = pd.read_excel(BytesIO(file_bytes), engine="openpyxl")

    parent_child_data = []
    columns = list(df.columns)

    for idx, row in df.iterrows():
        parts = []
        for col in columns:
            val = row[col]
            if pd.notna(val):
                parts.append(f"{col}: {val}")
        sentence = ", ".join(parts)
        if sentence.strip():
            parent_child_data.append({
                "parent_content": sentence,
                "metadata": {"source_type": "spreadsheet", "row_index": int(idx)},
                "children": [sentence]
            })

    return parent_child_data



# =============================================================================
# Main Processing Pipeline
# =============================================================================

async def process_document(
    file_bytes: bytes,
    file_name: str,
    mime_type: str,
    user_id: uuid.UUID,
    project_id: uuid.UUID,
    db: Session
) -> Document:
    """
    Pipeline xử lý tài liệu chính:
    1. Tính checksum → Checksum Guard
    2. Lưu file vật lý
    3. Tạo bản ghi Document
    4. Trích xuất văn bản + Chunking theo định dạng
    5. Sinh embedding qua Ollama BGE-M3
    6. Lưu chunks vào DB
    """
    checksum = compute_checksum(file_bytes)

    # Checksum Guard: chặn re-chunking trùng lặp
    existing = db.query(Document).filter(
        Document.project_id == project_id,
        Document.checksum == checksum
    ).first()
    if existing:
        raise ValueError(f"Tài liệu đã tồn tại (checksum trùng): {existing.file_name}")

    # Lưu file vật lý
    unique_name = f"{uuid.uuid4().hex}_{file_name}"
    storage_path = os.path.join(UPLOAD_DIR, unique_name)
    with open(storage_path, "wb") as f:
        f.write(file_bytes)

    # Tạo bản ghi Document
    doc = Document(
        id=uuid.uuid4(),
        user_id=user_id,
        project_id=project_id,
        source_type=_detect_source_type(file_name),
        file_name=file_name,
        mime_type=mime_type,
        storage_uri=storage_path,
        file_size=len(file_bytes),
        checksum=checksum,
        status=DocumentStatus.processing,
        created_at=datetime.now(timezone.utc),
        updated_at=datetime.now(timezone.utc)
    )
    db.add(doc)
    db.commit()

    try:
        # Chunking theo định dạng
        page_count = None
        ext = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else ""

        if ext == "txt":
            parent_child_data = chunk_plain_text_parent_child(file_bytes.decode("utf-8", errors="replace"))
        elif ext == "md":
            parent_child_data = chunk_markdown_parent_child(file_bytes.decode("utf-8", errors="replace"))
        elif ext == "docx":
            parent_child_data = chunk_docx_parent_child(file_bytes)
        elif ext == "pdf":
            parent_child_data, page_count = chunk_pdf_parent_child(file_bytes)
        elif ext in ("xlsx", "csv"):
            parent_child_data = chunk_excel_csv_parent_child(file_bytes, file_name)
        else:
            # Fallback: xử lý như plain text
            parent_child_data = chunk_plain_text_parent_child(file_bytes.decode("utf-8", errors="replace"))

        if page_count:
            doc.page_count = page_count

        if not parent_child_data:
            doc.status = DocumentStatus.failed
            doc.error_message = "Không trích xuất được nội dung từ tài liệu."
            db.commit()
            return doc

        # Thu thập tất cả child texts để sinh embedding
        all_child_texts = []
        for item in parent_child_data:
            all_child_texts.extend(item["children"])

        embeddings = []
        if all_child_texts:
            embeddings = await embedding_service.get_embeddings_batch(all_child_texts)

        # Lưu Parent & Child chunks vào DB
        child_emb_idx = 0
        chunk_counter = 0

        for item in parent_child_data:
            # 1. Lưu Parent Chunk (không có embedding, parent_id = None)
            parent_id = uuid.uuid4()
            db_parent = DocumentChunk(
                id=parent_id,
                project_id=project_id,
                document_id=doc.id,
                chunk_index=chunk_counter,
                content=item["parent_content"],
                token_count=count_tokens(item["parent_content"]),
                embedding=None,
                metadata_=item.get("metadata"),
                parent_id=None,
                created_at=datetime.now(timezone.utc),
                updated_at=datetime.now(timezone.utc)
            )
            db.add(db_parent)
            chunk_counter += 1

            # 2. Lưu Child Chunks (có embedding, liên kết tới parent)
            for child_text in item["children"]:
                emb = embeddings[child_emb_idx] if child_emb_idx < len(embeddings) else None
                child_emb_idx += 1

                db_child = DocumentChunk(
                    id=uuid.uuid4(),
                    project_id=project_id,
                    document_id=doc.id,
                    chunk_index=chunk_counter,
                    content=child_text,
                    token_count=count_tokens(child_text),
                    embedding=emb,
                    metadata_=item.get("metadata"),
                    parent_id=parent_id,
                    created_at=datetime.now(timezone.utc),
                    updated_at=datetime.now(timezone.utc)
                )
                db.add(db_child)
                chunk_counter += 1

        doc.status = DocumentStatus.completed
        db.commit()
        logger.info(f"Xử lý thành công tài liệu '{file_name}': {len(parent_child_data)} parent chunks và {len(all_child_texts)} child chunks.")

    except Exception as e:
        doc.status = DocumentStatus.failed
        doc.error_message = str(e)
        db.commit()
        logger.error(f"Lỗi xử lý tài liệu '{file_name}': {e}")
        raise

    return doc


def _detect_source_type(file_name: str) -> str:
    """Phát hiện loại nguồn từ tên file."""
    ext = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else "unknown"
    mapping = {
        "txt": "plain_text", "md": "markdown", "docx": "word",
        "pdf": "pdf", "xlsx": "excel", "csv": "csv"
    }
    return mapping.get(ext, "unknown")
